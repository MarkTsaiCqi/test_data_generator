import json
import csv
import xml.etree.ElementTree as ET
import os
from datetime import datetime
import random
from faker import Faker
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class FileGenerator:
    def __init__(self, output_dir="generated_data"):
        self.output_dir = output_dir
        self.fake = Faker()
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def generate_json(self):
        """生成 JSON 測試文件"""
        data = {
            "id": random.randint(1000, 9999),
            "name": self.fake.name(),
            "email": self.fake.email(),
            "address": {
                "street": self.fake.street_address(),
                "city": self.fake.city(),
                "country": self.fake.country()
            },
            "phone": self.fake.phone_number(),
            "created_at": datetime.now().isoformat()
        }
        
        filepath = os.path.join(self.output_dir, "test_data.json")
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def generate_txt(self):
        """生成 TXT 測試文件"""
        content = f"""測試文本文件
生成時間：{datetime.now()}
{self.fake.text(max_nb_chars=1000)}
"""
        filepath = os.path.join(self.output_dir, "test_data.txt")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

    def generate_md(self):
        """生成 Markdown 測試文件"""
        content = f"""# 測試文檔

## 基本信息
- 生成時間：{datetime.now()}
- 作者：{self.fake.name()}

## 內容概述
{self.fake.text(max_nb_chars=200)}

### 詳細信息
1. {self.fake.sentence()}
2. {self.fake.sentence()}
3. {self.fake.sentence()}

## 表格示例
| ID | 名稱 | 說明 |
|----|------|------|
| 1 | {self.fake.word()} | {self.fake.sentence()} |
| 2 | {self.fake.word()} | {self.fake.sentence()} |
"""
        filepath = os.path.join(self.output_dir, "test_data.md")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

    def generate_csv(self):
        """生成 CSV 測試文件"""
        headers = ['ID', '姓名', '郵箱', '電話', '地址']
        rows = []
        for i in range(10):
            rows.append([
                i + 1,
                self.fake.name(),
                self.fake.email(),
                self.fake.phone_number(),
                self.fake.address()
            ])
        
        filepath = os.path.join(self.output_dir, "test_data.csv")
        with open(filepath, 'w', encoding='utf-8', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(rows)

    def generate_xml(self):
        """生成 XML 測試文件"""
        root = ET.Element("TestData")
        
        users = ET.SubElement(root, "Users")
        for i in range(5):
            user = ET.SubElement(users, "User")
            ET.SubElement(user, "ID").text = str(i + 1)
            ET.SubElement(user, "Name").text = self.fake.name()
            ET.SubElement(user, "Email").text = self.fake.email()
            
        tree = ET.ElementTree(root)
        filepath = os.path.join(self.output_dir, "test_data.xml")
        tree.write(filepath, encoding='utf-8', xml_declaration=True)

    def generate_excel(self):
        """生成 Excel 測試文件（xlsx）"""
        # 創建數據
        data = {
            'ID': range(1, 11),
            '姓名': [self.fake.name() for _ in range(10)],
            '郵箱': [self.fake.email() for _ in range(10)],
            '電話': [self.fake.phone_number() for _ in range(10)],
            '地址': [self.fake.address() for _ in range(10)]
        }
        df = pd.DataFrame(data)
        
        # 保存為 xlsx
        xlsx_path = os.path.join(self.output_dir, "test_data.xlsx")
        df.to_excel(xlsx_path, index=False)

    def generate_doc(self):
        """生成 Word 文檔（docx 和 doc）"""
        doc = Document()
        
        # 添加標題
        doc.add_heading('測試文檔', 0)
        
        # 添加段落
        doc.add_paragraph(f'生成時間：{datetime.now()}')
        doc.add_paragraph(self.fake.text(max_nb_chars=500))
        
        # 添加表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'ID'
        header_cells[1].text = '名稱'
        header_cells[2].text = '說明'
        
        for i in range(3):
            cells = table.add_row().cells
            cells[0].text = str(i + 1)
            cells[1].text = self.fake.word()
            cells[2].text = self.fake.sentence()
        
        # 保存為 docx
        docx_path = os.path.join(self.output_dir, "test_data.docx")
        doc.save(docx_path)

    def generate_pdf(self):
        """生成 PDF 文件"""
        filepath = os.path.join(self.output_dir, "test_data.pdf")
        c = canvas.Canvas(filepath, pagesize=letter)
        
        # 添加標題
        c.setFont("Helvetica-Bold", 24)
        c.drawString(72, 750, "測試文檔")
        
        # 添加內容
        c.setFont("Helvetica", 12)
        c.drawString(72, 720, f"生成時間：{datetime.now()}")
        
        y = 700
        text = self.fake.text(max_nb_chars=500)
        for i in range(0, len(text), 80):
            c.drawString(72, y, text[i:i+80])
            y -= 20
        
        c.save()

def main():
    generator = FileGenerator()
    
    # 生成各種格式的文件
    generator.generate_json()
    generator.generate_txt()
    generator.generate_md()
    generator.generate_csv()
    generator.generate_xml()
    generator.generate_excel()
    generator.generate_doc()
    generator.generate_pdf()

if __name__ == "__main__":
    main() 