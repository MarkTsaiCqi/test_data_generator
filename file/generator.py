import os
import random
import string
from typing import Optional
import pandas as pd
import xlwt
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class FileGenerator:
    def __init__(self, output_dir: str = "generated_data"):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def _generate_random_text(self, length: int) -> str:
        """生成隨機文本"""
        return ''.join(random.choices(string.ascii_letters + string.digits + string.punctuation + ' ', k=length))

    def _generate_random_binary(self, length: int) -> bytes:
        """生成隨機二進制數據"""
        return bytes(random.getrandbits(8) for _ in range(length))

    def _generate_invalid_executable(self, filepath: str):
        """生成無效的執行檔"""
        # 生成一個小的無效二進制文件
        with open(filepath, 'wb') as f:
            f.write(self._generate_random_binary(1024))

    def _generate_invalid_video(self, filepath: str):
        """生成無效的影片文件"""
        # 生成一個無效的影片文件頭
        with open(filepath, 'wb') as f:
            # 寫入一些無效的影片文件頭數據
            f.write(b'\x00\x00\x00\x20\x66\x74\x79\x70')  # 模擬 MP4 文件頭
            f.write(self._generate_random_binary(1024))

    def _generate_code(self, language: str) -> str:
        """生成不同語言的示例代碼"""
        code_templates = {
            'c': '''
#include <stdio.h>

int main() {
    printf("Hello, World!\\n");
    return 0;
}
''',
            'cpp': '''
#include <iostream>

int main() {
    std::cout << "Hello, World!" << std::endl;
    return 0;
}
''',
            'cs': '''
using System;

class Program {
    static void Main() {
        Console.WriteLine("Hello, World!");
    }
}
''',
            'css': '''
body {
    font-family: Arial, sans-serif;
    margin: 0;
    padding: 20px;
}

.container {
    max-width: 800px;
    margin: 0 auto;
}
''',
            'go': '''
package main

import "fmt"

func main() {
    fmt.Println("Hello, World!")
}
''',
            'html': '''
<!DOCTYPE html>
<html>
<head>
    <title>Test Page</title>
</head>
<body>
    <h1>Hello, World!</h1>
</body>
</html>
''',
            'java': '''
public class Main {
    public static void main(String[] args) {
        System.out.println("Hello, World!");
    }
}
''',
            'js': '''
console.log("Hello, World!");

function example() {
    return "This is a test function";
}
''',
            'json': '''
{
    "name": "Test Data",
    "version": "1.0.0",
    "description": "Sample JSON file",
    "items": [
        {"id": 1, "value": "test1"},
        {"id": 2, "value": "test2"}
    ]
}
''',
            'md': '''
# Test Markdown File

This is a sample markdown file.

## Features
- Item 1
- Item 2
- Item 3

## Code Example
```python
print("Hello, World!")
```
''',
            'php': '''
<?php
echo "Hello, World!";
?>
''',
            'py': '''
def main():
    print("Hello, World!")

if __name__ == "__main__":
    main()
''',
            'rb': '''
def main
    puts "Hello, World!"
end

main
''',
            'sh': '''
#!/bin/bash
echo "Hello, World!"
''',
            'tex': '''
\\documentclass{article}
\\begin{document}
Hello, World!
\\end{document}
''',
            'ts': '''
function greet(name: string): string {
    return `Hello, ${name}!`;
}

console.log(greet("World"));
''',
            'txt': 'This is a sample text file.\nIt contains multiple lines of text.\nEach line demonstrates different content.',
            'csv': 'id,name,value\n1,test1,100\n2,test2,200\n3,test3,300',
            'xml': '''<?xml version="1.0" encoding="UTF-8"?>
<root>
    <item id="1">
        <name>test1</name>
        <value>100</value>
    </item>
    <item id="2">
        <name>test2</name>
        <value>200</value>
    </item>
</root>'''
        }
        return code_templates.get(language.lower(), '')

    def _generate_excel(self, filepath: str, format: str = 'xlsx'):
        """生成Excel文件"""
        df = pd.DataFrame({
            'id': [1, 2, 3],
            'name': ['test1', 'test2', 'test3'],
            'value': [100, 200, 300]
        })
        
        if format == 'xls':
            # 使用 xlwt 生成 .xls 文件
            workbook = xlwt.Workbook()
            sheet = workbook.add_sheet('Sheet1')
            
            # 寫入表頭
            headers = ['id', 'name', 'value']
            for col, header in enumerate(headers):
                sheet.write(0, col, header)
            
            # 寫入數據
            for row, (_, data) in enumerate(df.iterrows(), start=1):
                sheet.write(row, 0, int(data['id']))
                sheet.write(row, 1, data['name'])
                sheet.write(row, 2, int(data['value']))
            
            workbook.save(filepath)
        else:
            # 使用 pandas 生成 .xlsx 文件
            df.to_excel(filepath, index=False)

    def _generate_word(self, filepath: str, format: str = 'docx'):
        """生成Word文件"""
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a sample document.')
        doc.add_paragraph('It contains multiple paragraphs.')
        doc.save(filepath)

    def _generate_pdf(self, filepath: str):
        """生成PDF文件"""
        c = canvas.Canvas(filepath, pagesize=letter)
        c.drawString(100, 750, "Test PDF Document")
        c.drawString(100, 730, "This is a sample PDF file.")
        c.save()

    def generate_file(self, file_type: str, size_kb: Optional[int] = None) -> str:
        """生成指定類型的文件"""
        file_type = file_type.lower().lstrip('.')
        filename = f"test_file.{file_type}"
        filepath = os.path.join(self.output_dir, filename)

        if file_type in ['xlsx', 'xls']:
            self._generate_excel(filepath, file_type)
        elif file_type in ['docx', 'doc']:
            self._generate_word(filepath, file_type)
        elif file_type == 'pdf':
            self._generate_pdf(filepath)
        else:
            content = self._generate_code(file_type)
            if size_kb is not None:
                # 調整內容大小到指定大小
                target_size = size_kb * 1024
                while len(content.encode('utf-8')) < target_size:
                    content += self._generate_random_text(100)

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)

        return filepath

    def generate_invalid_file(self, file_type: str) -> str:
        """生成無效的測試文件"""
        file_type = file_type.lower().lstrip('.')
        filename = f"test_file.{file_type}"
        filepath = os.path.join(self.output_dir, filename)

        if file_type in ['exe', 'bat', 'apk', 'dll']:
            self._generate_invalid_executable(filepath)
        elif file_type in ['mp4', 'mov', 'avi', 'mkv']:
            self._generate_invalid_video(filepath)
        else:
            raise ValueError(f"不支持的無效文件類型: {file_type}")

        return filepath

    def generate_all(self) -> list:
        """生成所有支持的文件類型"""
        file_types = [
            # 原有文件類型
            'json', 'txt', 'md', 'csv', 'xml', 'xlsx', 'xls', 'docx', 'doc', 'pdf',
            # 新增文件類型
            'c', 'cpp', 'cs', 'css', 'go', 'html', 'java', 'js', 'php', 'py', 
            'rb', 'sh', 'tex', 'ts'
        ]
        generated_files = []
        
        for file_type in file_types:
            try:
                filepath = self.generate_file(file_type)
                generated_files.append(filepath)
                print(f"Generated {file_type.upper()} file: {filepath}")
            except Exception as e:
                print(f"Error generating {file_type.upper()} file: {str(e)}")
        
        return generated_files

    def generate_all_invalid(self) -> list:
        """生成所有無效的測試文件"""
        file_types = [
            # 執行檔
            'exe', 'bat', 'apk', 'dll',
            # 影片檔
            'mp4', 'mov', 'avi', 'mkv'
        ]
        generated_files = []
        
        for file_type in file_types:
            try:
                filepath = self.generate_invalid_file(file_type)
                generated_files.append(filepath)
                print(f"Generated invalid {file_type.upper()} file: {filepath}")
            except Exception as e:
                print(f"Error generating invalid {file_type.upper()} file: {str(e)}")
        
        return generated_files

def create_pdf_from_text_file(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as f:
        text_content = f.read()

    # 創建 PDF
    c = canvas.Canvas(output_file, pagesize=letter)
    width, height = letter
    c.setFont("Helvetica", 12)
    text_object = c.beginText(40, height - 40)

    for line in text_content.splitlines():
        text_object.textLine(line)
    c.drawText(text_object)
    c.save()
    print(f'PDF 檔案已生成：{output_file}')

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("請提供輸入檔案和輸出檔案名稱。")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    create_pdf_from_text_file(input_file, output_file) 